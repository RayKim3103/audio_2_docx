from __future__ import annotations

import zipfile
from pathlib import Path

from .paths import PANDOC_DIR, REFERENCE_DIR, configure_environment
from .utils import clean_text_for_xml, read_text, write_text


def set_east_asia_font(run, font_name: str) -> None:
    try:
        from docx.oxml.ns import qn
        run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    except Exception:
        pass


def create_reference_docx(font_name: str = "Malgun Gothic", font_size_pt: int = 8) -> Path:
    configure_environment()
    ref = REFERENCE_DIR / f"reference_{font_size_pt}pt.docx"
    if ref.exists():
        return ref
    from docx import Document
    from docx.shared import Pt
    doc = Document()
    styles = doc.styles
    for style_name in ["Normal", "Body Text", "List Paragraph", "Table Grid", "Heading 1", "Heading 2", "Heading 3"]:
        try:
            style = styles[style_name]
            style.font.name = font_name
            style.font.size = Pt(font_size_pt)
            if style_name.startswith("Heading"):
                style.font.bold = True
        except Exception:
            pass
    p = doc.add_paragraph("reference")
    if p.runs:
        p.runs[0].font.size = Pt(font_size_pt)
        p.runs[0].font.name = font_name
        set_east_asia_font(p.runs[0], font_name)
    doc.save(ref)
    return ref


def ensure_pandoc() -> None:
    configure_environment()
    import pypandoc
    try:
        pypandoc.get_pandoc_path()
        return
    except Exception:
        pass
    try:
        pypandoc.download_pandoc(targetfolder=str(PANDOC_DIR))
    except Exception as e:
        raise RuntimeError(
            "Pandoc을 찾거나 다운로드하지 못했습니다. pypandoc-binary 설치 상태 또는 네트워크를 확인하세요."
        ) from e


def force_docx_font_size(docx_path: Path, font_name: str = "Malgun Gothic", font_size_pt: int = 8) -> None:
    from docx import Document
    from docx.shared import Pt
    doc = Document(docx_path)
    for style in doc.styles:
        try:
            style.font.name = font_name
            style.font.size = Pt(font_size_pt)
        except Exception:
            pass
    def apply_para(p):
        for run in p.runs:
            run.font.name = font_name
            run.font.size = Pt(font_size_pt)
            set_east_asia_font(run, font_name)
    for p in doc.paragraphs:
        apply_para(p)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    apply_para(p)
    for sec in doc.sections:
        for p in sec.header.paragraphs:
            apply_para(p)
        for p in sec.footer.paragraphs:
            apply_para(p)
    doc.save(docx_path)


def validate_docx(docx_path: Path) -> None:
    with zipfile.ZipFile(docx_path, "r") as zf:
        bad = zf.testzip()
        if bad:
            raise RuntimeError(f"DOCX zip 구조 오류: {bad}")
        if "word/document.xml" not in zf.namelist():
            raise RuntimeError("DOCX 내부에 word/document.xml이 없습니다.")
    from docx import Document
    Document(docx_path)


def markdown_to_docx(md_path: Path, docx_path: Path, font_size_pt: int = 8, font_name: str = "Malgun Gothic") -> Path:
    configure_environment()
    md_path = Path(md_path)
    docx_path = Path(docx_path)
    docx_path.parent.mkdir(parents=True, exist_ok=True)
    # Sanitize invalid XML characters before pandoc conversion.
    safe_md = md_path.with_suffix(".safe.md")
    write_text(safe_md, clean_text_for_xml(read_text(md_path)))
    ensure_pandoc()
    import pypandoc
    ref = create_reference_docx(font_name=font_name, font_size_pt=font_size_pt)
    extra_args = ["--standalone", "--reference-doc", str(ref)]
    pypandoc.convert_file(str(safe_md), "docx", outputfile=str(docx_path), extra_args=extra_args)
    force_docx_font_size(docx_path, font_name=font_name, font_size_pt=font_size_pt)
    validate_docx(docx_path)
    return docx_path
